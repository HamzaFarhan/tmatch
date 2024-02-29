import re
from pathlib import Path
from typing import Callable

import pandas as pd
from docx import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import (
    PyPDFLoader,
    UnstructuredWordDocumentLoader,
)
from termcolor import colored

from .utils import clean_text, resolve_data_path

CHUNK_SIZE = None
CHUNK_OVERLAP = 0
MAX_CHUNK_SIZE = 100_000
EXTENSIONS = [".pdf", ".doc", ".docx", ".txt"]


def create_path_df(
    data_path: Path | str,
    extensions: list[str] = EXTENSIONS,
    extra_cols: dict[str, str | Path] = {},
) -> pd.DataFrame:
    if Path(data_path).is_dir():
        paths = [
            str(path)
            for path in resolve_data_path(data_path)
            if Path(path).suffix in extensions
        ]
    else:
        paths = [str(data_path)]
    df_dict = {"path": paths}
    df_dict.update(extra_cols)
    return pd.DataFrame(df_dict)


def load_and_split_resume_text(
    loader: Callable | None = None,
    chunk_size: int | None = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
    text: str | list[str] | None = None,
    map_fns: list[Callable] | None = None,
) -> list[str]:
    assert loader or text, "Either loader or text must be provided"
    map_fns = map_fns or []
    text_splitter = None
    if chunk_size is not None:
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=min(chunk_size, MAX_CHUNK_SIZE),
            chunk_overlap=chunk_overlap,
            separators=["\n\n", " ", ".", ""],
            keep_separator=False,
        )
    if loader:
        if text_splitter is None:
            texts = loader.load()
        else:
            texts = loader.load_and_split(text_splitter)
        map_fns.insert(0, lambda text: text.page_content)
    elif text:
        if isinstance(text, list):
            text = " ".join(text)
        if text_splitter is None:
            texts = [text]
        else:
            texts = text_splitter.split_text(text)
    else:
        return []
    map_fns += [clean_text]
    ret_texts = []
    # print(f"\nMAP FNS: {map_fns}\n")
    for text in texts:
        # print(f"\nTEXT: {text}\n")
        for map_fn in map_fns:
            text = map_fn(text)
        if text.strip() != "":
            ret_texts.append(text)
    return ret_texts


def extract_txt_resume_text(
    file: str | Path,
    chunk_size: int | None = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
) -> list[str]:
    file = str(file)
    print(colored(f"\n\nExtracting text from TXT file: {file}\n\n", color="cyan"))
    with open(str(file), "r") as f:
        text = f.read().splitlines()
    return load_and_split_resume_text(
        text=text, chunk_size=chunk_size, chunk_overlap=chunk_overlap
    )


def extract_pdf_resume_text(
    file: str | Path,
    chunk_size: int | None = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
) -> list[str]:
    file = str(file)
    print(colored(f"\n\nExtracting text from PDF file: {file}\n\n", color="cyan"))
    pdf_loader = PyPDFLoader(str(file))
    return load_and_split_resume_text(
        loader=pdf_loader, chunk_size=chunk_size, chunk_overlap=chunk_overlap
    )


def extract_docx_resume_text(
    file: str | Path,
    chunk_size: int | None = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
) -> list[str]:
    file = str(file)
    print(colored(f"\n\nExtracting text from WORD file: {file}\n\n", color="cyan"))
    docx_loader = UnstructuredWordDocumentLoader(str(file))
    return load_and_split_resume_text(
        loader=docx_loader, chunk_size=chunk_size, chunk_overlap=chunk_overlap
    )


def extract_docx_resume_text_2(
    file: str | Path,
    chunk_size: int | None = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
) -> list[str]:
    file = str(file)
    print(colored(f"\n\nExtracting text from WORD file: {file}\n\n", color="cyan"))
    doc = Document(file)
    text = [para.text.strip() for para in doc.paragraphs]

    table_text = []
    for table in doc.tables:
        for col in table.columns:
            for cell in col.cells:
                txt = cell.text.strip()
                if len(txt) > 1:
                    table_text.append(txt)
    text += table_text
    return load_and_split_resume_text(
        text=text,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        map_fns=[lambda txt: re.sub(r" +", " ", txt)],
    )


def extract_resume_text(
    file: str | Path,
    chunk_size: int | None = CHUNK_SIZE,
    chunk_overlap: int = CHUNK_OVERLAP,
) -> list[str]:
    suff = Path(file).suffix
    suff_extractor_map = {
        ".pdf": [extract_pdf_resume_text],
        ".doc": [extract_docx_resume_text, extract_docx_resume_text_2],
        ".docx": [extract_docx_resume_text, extract_docx_resume_text_2],
        ".txt": [extract_txt_resume_text],
    }
    extractors = suff_extractor_map.get(suff, None)
    if not extractors:
        print(colored(f"\n\nNo extractor for {suff} file: {file}\n\n", color="red"))
        return []
    for extractor in extractors:
        try:
            return extractor(
                file=file, chunk_size=chunk_size, chunk_overlap=chunk_overlap
            )
        except Exception as e:
            print(
                colored(
                    f"\n\nCould not extract text from {suff} file: {file} using extractor: {extractor.__name__}\n\n\t{e}\n\n",
                    color="red",
                )
            )
            continue
    return []
